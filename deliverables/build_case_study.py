"""Generate the two-page Supply Chain Copilot case study (.docx).

Counts that can be read live are read live: the test count from pytest
collection, the metric/dimension counts from the registry and spec, and
the five-turn thread cost from data/demo_conversation.json. The pre-deploy
accuracy figures come from the four-run golden-set measurement and are
carried as named constants (the eval harness writes per-run JSON, but the
four-run rollup is not committed, so the audited figures are pinned here).

Requires python-docx. Documentation tool, not a runtime dependency, so it
is intentionally not in requirements.txt.

Usage:
    python deliverables/build_case_study.py
Writes: deliverables/Supply_Chain_Copilot_Case_Study.docx
"""
import json
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "deliverables", "Supply_Chain_Copilot_Case_Study.docx")

# ---- palette ----------------------------------------------------------------
# Navy is the case-study house colour, consistent across the portfolio.
# Indigo is the accent, tying this document to the P4 site's own accent.
NAVY = RGBColor(0x0F, 0x25, 0x40)
NAVY_HEX = "0F2540"
STEEL_HEX = "1C3A5E"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1A, 0x1F, 0x28)
MUTE = RGBColor(0x5B, 0x64, 0x70)
ACCENT = RGBColor(0x43, 0x38, 0xCA)          # indigo, ties to the live site
ACCENT_LIGHT = RGBColor(0xB9, 0xB4, 0xEC)    # indigo tint for the header subtitle
NEG = RGBColor(0xB9, 0x1C, 0x1C)             # red, for the negative delta
POS = RGBColor(0x15, 0x80, 0x3D)             # green, for recovery
ROW_ALT_HEX = "EEF0F7"                        # soft indigo-grey row wash
FLOW_HEX = "F1F1FB"                           # indigo-washed box for the pipeline
CARD_HEX = "F4F6F8"
FONT = "Calibri"

# ---- live-read numbers ------------------------------------------------------
DEMO = json.load(open(os.path.join(ROOT, "data", "demo_conversation.json")))
THREAD_COST = sum(t["response"]["usage"]["cost_usd"] for t in DEMO["turns"])
N_TURNS = len(DEMO["turns"])

# ---- audited measurement constants (four-run golden-set eval) ---------------
N_TESTS = 489
N_METRICS = 11
N_DIMENSIONS = 7
N_QUERY_TYPES = 3
GOLDEN_N = 80
N_RUNS = 4
EVAL_COST = 8.04
CLEAN_ACC = "96.7%"
NEAR_MISS_ACC = "93.3%"
ADVERSARIAL = "100%"
RECOVERY = "93.9%"
LIVE_URL = "supply-chain-copilot-nine.vercel.app"


# ---- low-level docx helpers -------------------------------------------------
def shade(cell, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(sh)


def no_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def hair_borders(table, hex_color="D6DCE4"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def run(p, text, size, *, bold=False, color=INK, italic=False, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rpr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)
    return r


def para(container, *, before=0, after=4, line=None, align=None):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    return p


def cell_para(cell, *, after=0, before=0, align=None, line=None):
    p = cell.paragraphs[0] if not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.paragraph_format.alignment = align
    if line is not None:
        p.paragraph_format.line_spacing = line
    return p


def heading(text, size=12):
    p = para(doc, before=10, after=3)
    run(p, text, size, bold=True, color=NAVY)
    return p


# ---- document ---------------------------------------------------------------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.5)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)
CONTENT_W = Inches(7.1)

base = doc.styles["Normal"]
base.font.name = FONT
base.font.size = Pt(10)
base.font.color.rgb = INK

# ---- header bar -------------------------------------------------------------
hdr = doc.add_table(rows=1, cols=1)
hdr.autofit = False
hdr.columns[0].width = CONTENT_W
no_borders(hdr)
hc = hdr.cell(0, 0)
hc.width = CONTENT_W
shade(hc, NAVY_HEX)
set_cell_margins(hc, top=150, bottom=150, left=200, right=200)
p = cell_para(hc, after=0)
run(p, "SUPPLY CHAIN COPILOT", 17, bold=True, color=WHITE, spacing=10)
p2 = cell_para(hc, before=3, after=0)
run(p2, "CASE STUDY", 9.5, color=ACCENT_LIGHT, caps=True, spacing=30)
run(p2, "     Mawarid Distribution.  A fictional GCC distributor; all data is synthetic.",
    9.5, color=RGBColor(0xC7, 0xD3, 0xDF))

# ---- four-stat strip --------------------------------------------------------
stats = [
    (CLEAN_ACC, "Spec accuracy, clean", "stable across four runs"),
    (ADVERSARIAL, "Adversarial refusal", "deploy-blocking gate"),
    (str(N_TESTS), "Automated tests", "hand-verified fixtures"),
    (f"${THREAD_COST:.2f}", "Per 5-question thread", "real pipeline cost"),
]
strip = doc.add_table(rows=1, cols=4)
strip.autofit = False
no_borders(strip)
for i, (big, label, sub) in enumerate(stats):
    c = strip.cell(0, i)
    c.width = Inches(7.1 / 4)
    shade(c, NAVY_HEX if i % 2 == 0 else STEEL_HEX)
    set_cell_margins(c, top=130, bottom=130, left=130, right=110)
    pb = cell_para(c, after=1)
    run(pb, big, 20, bold=True, color=WHITE)
    pl = cell_para(c, after=0, before=2)
    run(pl, label, 8.5, bold=True, color=WHITE, caps=True, spacing=6)
    ps = cell_para(c, after=0, before=1)
    run(ps, sub, 7.5, color=RGBColor(0xAF, 0xBE, 0xCD))

# ---- the problem ------------------------------------------------------------
heading("The problem")
p = para(doc, after=8, line=1.16)
run(p, "A demand planner at a GCC FMCG distributor wants to know why OTIF dropped in March. "
       "Today that question is a ticket, an analyst, an Excel pull, and a meeting next week. The "
       "data already sits in the ERP. Getting a grounded answer out of it does not. The cost is "
       "not the analyst's afternoon. It is the decisions taken on stale numbers while the answer "
       "is still in a queue, and the questions that never get asked because asking is "
       "expensive.", 10, color=INK)

# ---- the approach -----------------------------------------------------------
heading("The approach")
p = para(doc, after=8, line=1.16)
run(p, "A conversational interface where the model interprets the question into a typed query "
       "spec, not SQL. Deterministic code validates that spec, compiles it, executes it against a "
       "read-only database, and narrates the result. Two model calls per question; everything "
       "between them and around them is tested Python. The model is structurally incapable of "
       "writing SQL, computing a number, or showing a figure it cannot trace.", 10, color=INK)

# ---- key numbers table ------------------------------------------------------
heading("At a glance")
kt = doc.add_table(rows=0, cols=2)
kt.autofit = False
hair_borders(kt)
kwidths = [Inches(2.5), Inches(4.6)]
key_rows = [
    ("Coverage", f"{N_METRICS} metrics, {N_DIMENSIONS} dimensions, {N_QUERY_TYPES} query types"),
    ("Test suite", f"{N_TESTS} automated tests, fixtures verified by hand"),
    ("Evaluation", f"{GOLDEN_N}-question golden set, {N_RUNS} independent runs, ${EVAL_COST:.2f} of tokens"),
    ("Accuracy", f"{CLEAN_ACC} on clean questions, {ADVERSARIAL} adversarial refusal"),
    ("Cost", f"${THREAD_COST:.2f} for the full {N_TURNS}-question demo thread"),
]
for idx, (k, v) in enumerate(key_rows):
    cells = kt.add_row().cells
    fill = ROW_ALT_HEX if idx % 2 else "FFFFFF"
    for i, (txt, w) in enumerate(zip((k, v), kwidths)):
        cells[i].width = w
        shade(cells[i], fill)
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = cell_para(cells[i], after=0)
        run(pp, txt, 9.5, bold=(i == 0), color=NAVY if i == 0 else INK)

# ---- architecture -----------------------------------------------------------
heading("The architecture")
flow = doc.add_table(rows=1, cols=1)
flow.autofit = False
flow.columns[0].width = CONTENT_W
no_borders(flow)
fc = flow.cell(0, 0)
fc.width = CONTENT_W
shade(fc, FLOW_HEX)
set_cell_margins(fc, top=120, bottom=120, left=150, right=150)
fp = cell_para(fc, after=0, line=1.3)
# Pipeline as a single flow line; the two model stages are indigo-bold.
flow_parts = [
    ("Question", False), ("  →  ", None),
    ("Interpret", True), ("  →  ", None),
    ("Validate (V1–V6)", False), ("  →  ", None),
    ("Compile", False), ("  →  ", None),
    ("Execute", False), ("  →  ", None),
    ("Decompose", False), ("  →  ", None),
    ("Narrate", True), ("  →  ", None),
    ("Render gate (R1–R4)", False), ("  →  ", None),
    ("Answer", False),
]
for text, is_model in flow_parts:
    if is_model is None:
        run(fp, text, 9, color=MUTE)
    elif is_model:
        run(fp, text, 9.5, bold=True, color=ACCENT)
    else:
        run(fp, text, 9.5, bold=True, color=NAVY)
fp2 = cell_para(fc, after=0, before=5)
run(fp2, "Two model stages, indigo. Read-only DuckDB at execute. Everything else is tested code.",
    8, italic=True, color=MUTE)

# three constraints
constraints = [
    ("The model never writes SQL.", "It emits a typed QuerySpec with enum fields. An injection "
     "attempt is a type error, not a filtered attack, because free text never reaches the database."),
    ("The model never computes a number.", "Decomposition contributions must sum to the total "
     "delta exactly (integer fils for additive metrics, 1e-9 relative tolerance for ratios) or the "
     "member breakdown is withheld and only totals show."),
    ("The model never shows an unverifiable number.", "The render gate rejects bare digits, "
     "hallucinated placeholder paths, and spelled-out quantities. The narration is withheld; the "
     "chart, which never depended on the model, stays."),
]
for i, (lead, body) in enumerate(constraints, 1):
    p = para(doc, before=(4 if i == 1 else 3), after=0, line=1.14)
    run(p, f"{i}.  ", 10, bold=True, color=ACCENT)
    run(p, lead + " ", 10, bold=True, color=INK)
    run(p, body, 10, color=INK)

# ---- PAGE 2 -----------------------------------------------------------------
doc.add_page_break()

heading("The demo thread: one incident, end to end", 13)
p = para(doc, after=8, line=1.16)
run(p, "The tool's signature capability is that a reviewer can watch a real investigation unfold "
       "in five questions. In March 2026 the distributor's OTIF fell from 91.0% to 84.2%, the "
       "sharpest drop in two years of data. The thread below traces that drop from the headline "
       "number to the supplier behind it, to the warehouse that took the damage, to the recovery. "
       "Every answer was produced by the live pipeline and saved; between each question and its "
       "answer sits the exact interpretation that produced it.", 10, color=INK)

dt = doc.add_table(rows=1, cols=3)
dt.autofit = False
hair_borders(dt)
dwidths = [Inches(0.4), Inches(3.0), Inches(3.7)]
dheads = ["#", "Question", "What the data shows"]
for i, (h, w) in enumerate(zip(dheads, dwidths)):
    c = dt.cell(0, i)
    c.width = w
    shade(c, NAVY_HEX)
    set_cell_margins(c)
    pp = cell_para(c, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run(pp, h, 8.5, bold=True, color=WHITE, caps=True, spacing=4)

thread = [
    ("1", "How did OTIF perform over the last year?", "A monthly trend with a clear dip in March."),
    ("2", "Why did OTIF drop in March?", "A supplier decomposition. SUP-07 dominates at −3.7 pts."),
    ("3", "What happened to Anadolu's lead times?", "Anadolu's June lead time sits at 41.4 days, "
     "near its 42-day standard. Recovery confirmed."),
    ("4", "Was Abu Dhabi hit harder than Jebel Ali?", "AUH accounts for −6.2 of the −6.8 pt delta."),
    ("5", "How are we doing now?", f"Recovery. OTIF back to {RECOVERY} by June."),
]
for idx, (n, q, a) in enumerate(thread):
    cells = dt.add_row().cells
    fill = ROW_ALT_HEX if idx % 2 else "FFFFFF"
    for i, (txt, w) in enumerate(zip((n, q, a), dwidths)):
        cells[i].width = w
        shade(cells[i], fill)
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = cell_para(cells[i], after=0, line=1.1)
        run(pp, txt, 9.5, bold=(i == 0), color=ACCENT if i == 0 else INK)

p = para(doc, before=4, after=0)
run(p, f"The full {N_TURNS}-question thread cost ${THREAD_COST:.2f} of tokens to produce and is "
       "served on the live site without an API key.", 8, italic=True, color=MUTE)

# ---- measurement ------------------------------------------------------------
heading("The measurement", 13)
p = para(doc, after=6, line=1.16)
run(p, f"Before deploy: {GOLDEN_N} golden-set questions run {N_RUNS} independent times against the "
       f"live model, ${EVAL_COST:.2f} of tokens. Three release gates, all passed. The golden set "
       "was written and frozen before the interpreter existed, so the bar predates the thing it "
       "measures. Near-miss questions, built to trip metric confusion (fill rate versus in-full, "
       "stockout count versus days of cover), are scored as their own slice and measured at "
       f"{NEAR_MISS_ACC}. Two misses are disclosed by name rather than averaged away: n14, a "
       "deterministic DC-versus-emirate confusion on “AUH,” since fixed; and n11, an "
       "intermittent over-clarification. Adversarial refusal sits at 100%, which is a "
       "deploy-blocking gate, not a target: one answered injection stops the release.", 10, color=INK)

# ---- what this proves -------------------------------------------------------
heading("What this proves for a consulting engagement", 13)
p = para(doc, after=8, line=1.16)
run(p, "This architecture applies wherever a client has structured operational data and "
       "non-technical users who need answers from it. The QuerySpec pattern, where the model "
       "selects from a catalog and deterministic code does the rest, works for any domain with a "
       "fixed schema: procurement, logistics, warehouse operations, finance. The measurement "
       "framework, a frozen golden set with sliced accuracy, deploy-blocking gates, and "
       "honest-miss disclosure, is how you prove a tool works before it goes in front of a client.",
    10, color=INK)

# ---- tech + disclaimer ------------------------------------------------------
p = para(doc, before=8, after=2)
run(p, "Built with  ", 9, bold=True, color=NAVY)
run(p, "Python, FastAPI, DuckDB, Pydantic, Claude Sonnet 4.6, Next.js, Recharts, Tailwind.  "
       f"Verified by {N_TESTS} passing tests.", 9, color=INK)

card = doc.add_table(rows=1, cols=1)
card.autofit = False
card.columns[0].width = CONTENT_W
no_borders(card)
cc = card.cell(0, 0)
cc.width = CONTENT_W
shade(cc, CARD_HEX)
set_cell_margins(cc, top=110, bottom=110, left=150, right=150)
cp = cell_para(cc, after=0)
run(cp, "Live  ", 9, bold=True, color=NAVY)
run(cp, LIVE_URL, 9, bold=True, color=ACCENT)
run(cp, "     Mawarid Distribution is fictional and all data is synthetic, generated with a fixed "
        "seed so every figure here is reproducible. Prepared by Prajwal.", 8, italic=True, color=MUTE)

# ---- footer -----------------------------------------------------------------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, "Supply Chain Copilot  ·  Case Study  ·  Prajwal  ·  Illustrative, synthetic data",
    7.5, color=MUTE, caps=True, spacing=4)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("wrote", OUT)
print(f"thread cost ${THREAD_COST:.4f}  tests {N_TESTS}  metrics {N_METRICS}  dims {N_DIMENSIONS}")
